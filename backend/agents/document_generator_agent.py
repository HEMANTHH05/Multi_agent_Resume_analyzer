import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .base_agent import BaseAgent


def add_horizontal_rule(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4F46E5")
    pBdr.append(bottom)
    pPr.append(pBdr)


class DocumentGeneratorAgent(BaseAgent):
    name = "DocumentGeneratorAgent"
    step_index = 8

    async def run(self) -> dict:
        final_resume = self.context.get("final_resume", {})
        session_id = self.context.get("session_id", "resume")

        await self.stream_to_client("Generating your Word document...\n")

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        # --- Name & Contact (from original resume text) ---
        resume_text = self.context.get("resume_text", "")
        name_line = resume_text.split("\n")[0].strip() if resume_text else "Candidate Name"

        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name_line)
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)

        # --- Summary ---
        summary = final_resume.get("summary", "")
        if summary:
            self._add_section_header(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph(summary)
            summary_para.style.font.size = Pt(10)
            self._set_para_spacing(summary_para)

        # --- Skills ---
        skills = final_resume.get("skills_section", [])
        if skills:
            self._add_section_header(doc, "TECHNICAL SKILLS")
            skills_text = " • ".join(skills)
            skills_para = doc.add_paragraph(skills_text)
            skills_para.style.font.size = Pt(10)
            self._set_para_spacing(skills_para)

        # --- Experience ---
        experiences = final_resume.get("experiences", [])
        if experiences:
            self._add_section_header(doc, "PROFESSIONAL EXPERIENCE")

        for exp in experiences:
            company = exp.get("company", "")
            title = exp.get("title", "")
            dates = exp.get("dates", "")

            # Company + dates line
            co_para = doc.add_paragraph()
            co_para.paragraph_format.space_before = Pt(8)
            co_para.paragraph_format.space_after = Pt(0)
            co_run = co_para.add_run(company)
            co_run.bold = True
            co_run.font.size = Pt(11)
            co_run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)

            if dates:
                tab_stop = OxmlElement("w:tab")
                co_para._p.append(tab_stop)
                dates_run = co_para.add_run(f"\t{dates}")
                dates_run.font.size = Pt(10)
                dates_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

            # Title line
            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(4)
            title_run = title_para.add_run(title)
            title_run.italic = True
            title_run.font.size = Pt(10)
            title_run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

            # Bullets
            for bullet in exp.get("bullets", []):
                bullet_para = doc.add_paragraph(style="List Bullet")
                bullet_para.paragraph_format.space_before = Pt(1)
                bullet_para.paragraph_format.space_after = Pt(1)
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                bullet_run = bullet_para.add_run(bullet)
                bullet_run.font.size = Pt(10)

        # Save
        output_dir = "/tmp"
        os.makedirs(output_dir, exist_ok=True)
        docx_path = os.path.join(output_dir, f"{session_id}_resume.docx")
        doc.save(docx_path)

        await self.stream_to_client(f"✅ Resume document generated successfully!\n")

        return {"docx_path": docx_path}

    def _add_section_header(self, doc: Document, title: str):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        add_horizontal_rule(para)

    def _set_para_spacing(self, para):
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(6)
